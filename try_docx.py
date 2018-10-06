import json

from docx import Document
from docx.shared import Cm, Pt

from core import settings


def get_settings():
    return {
        "font-family": "Times New Roman",
        "font-size": 12,
        "line-spacing": 1.15,
        "margin": [1, 1, 1, 1],
    }


def insert_table(doc, data=None, columns_settings=None):
    def get_data_stub():
        with open("data/output.txt", "r", encoding="utf-8") as file:
            return json.loads(file.read())

    def columns_settings_stub():
        return {
            "syn": True,
            "ant": True,
            "def": True,
        }

    def compose_header(cols_settings):
        header_row = ["№\nп/п", "Слово", "Перевод"]
        header_composer = [
            ("syn", "Синонимы"),
            ("ant", "Антонимы"),
            ("def", "Определения")
        ]

        for key, item in header_composer:
            if cols_settings.get(key, False):
                header_row.append(item)

        return header_row

    def insert_bold_category(cell, category):
        if cell.paragraphs[0].text != "":
            cell.add_paragraph()

        to_be_bold = cell.paragraphs[len(cell.paragraphs) - 1].add_run(category)
        to_be_bold.font.bold = True

    def insert_lexical_category(cell, category, key, constraint=None):
        insert_bold_category(cell, category["lexicalCategory"])

        for item in category.get(key, [])[:constraint]:
            cell.add_paragraph("- " + item + ";")

    def insert_header(table, header):
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = header[i]

        table.add_row()

    data = get_data_stub()
    columns_settings = columns_settings_stub()
    header = compose_header(columns_settings)

    table = doc.add_table(rows=len(data.keys()), cols=len(header))
    table.style = 'Table Grid'
    insert_header(table, header)

    i = 0

    for word, word_data in data.items():
        i += 1
        row = table.rows[i]
        table_builder = (
            (3, "syn", "synonyms", settings.user["output"]["synonyms_count"]),
            (4, "ant", "antonyms", settings.user["output"]["antonyms_count"]),
            (5, "def", "definitions", None),
        )

        row.cells[0].text = str(i)
        row.cells[1].text = word

        for lexical_category in word_data["translations"]["lexicalCategories"]:
            insert_lexical_category(row.cells[2], lexical_category, "translations")

        for lexical_category in word_data["api"][0]["lexicalEntries"]:
            for idx, settings_key, cat_key, settings_constraint in table_builder:
                if columns_settings.get(settings_key, False):
                    insert_lexical_category(row.cells[idx], lexical_category, cat_key, settings_constraint)


def switch_orientation(doc):
    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = not section.orientation
    section.page_width = new_width
    section.page_height = new_height


def set_margins(doc, margin):
    section = doc.sections[-1]
    section.top_margin = Cm(margin[0])
    section.right_margin = Cm(margin[1])
    section.bottom_margin = Cm(margin[2])
    section.left_margin = Cm(margin[3])


def change_font(doc, family, size, interval):
    style = doc.styles['Normal']
    style.font.name = family
    style.font.size = Pt(size)
    style.paragraph_format.line_spacing = interval


if __name__ == "__main__":
    document = Document()
    user_settings = get_settings()

    # TODO: уточнить
    switch_orientation(document)

    set_margins(document, user_settings["margin"])
    change_font(document, user_settings["font-family"], user_settings["font-size"], user_settings["line-spacing"])

    insert_table(document)

    # records = stub_rows()
    # records = get_records()
    #
    # table = document.add_table(rows=len(records), cols=6)
    # table.style = 'Table Grid'
    #
    # for i, row in enumerate(table.rows):
    #     for j, cell in enumerate(row.cells):
    #         cell.text = records[i][j]

    document.save('demo3.docx')