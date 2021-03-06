import json

from docx import Document
from docx.shared import Cm, Pt

from core import settings
from .abstracts import ResultDataWriter


class TxtDataWriter(ResultDataWriter):
    def __init__(self, path):
        self.path = path

    def write(self, data):
        with open(self.path, "w", encoding="utf-8") as file:
            file.write(json.dumps(data, ensure_ascii=False))


class DocDataWriter(ResultDataWriter):
    def __init__(self, filename, columns_settings):
        self.filename = filename
        self.settings = settings.user["docx-writer"]
        self.columns_settings = columns_settings

    def write(self, data):
        def switch_orientation(doc):
            section = doc.sections[-1]
            new_width, new_height = section.page_height, section.page_width
            section.orientation = not section.orientation
            section.page_width = new_width
            section.page_height = new_height

        def set_margins(doc, margin):
            section = doc.sections[-1]
            section.top_margin = Cm(margin[0])
            section.right_margin = Cm(margin[1])
            section.bottom_margin = Cm(margin[2])
            section.left_margin = Cm(margin[3])

        def change_font(doc, family, size, interval):
            style = doc.styles['Normal']
            style.font.name = family
            style.font.size = Pt(int(size))
            style.paragraph_format.line_spacing = float(interval)

        def compose_header(cols_settings):
            header_row = ["№\nп/п", "Слово", "Перевод"]
            header_composer = [
                ("syn", "Синонимы"),
                ("ant", "Антонимы"),
                ("def", "Определения")
            ]

            for key, item in header_composer:
                if cols_settings.get(key, False):
                    header_row.append(item)

            return header_row

        def insert_lexical_category(cell, category, key, constraint=None):
            if cell.paragraphs[0].text != "":
                cell.add_paragraph()

            to_be_bold = cell.paragraphs[len(cell.paragraphs) - 1].add_run(category["lexicalCategory"])
            to_be_bold.font.bold = True

            for item in category.get(key, [])[:constraint]:
                cell.add_paragraph("- " + item + ";")

        def insert_header(doc_table, table_header):
            doc_table.add_row()
            for i, cell in enumerate(doc_table.rows[0].cells):
                cell.text = table_header[i]

        document = Document()
        switch_orientation(document)
        set_margins(document, self.settings["margin"])
        change_font(document, self.settings["font-family"], self.settings["font-size"], self.settings["line-spacing"])

        header = compose_header(self.columns_settings)
        table = document.add_table(rows=len(data.keys()), cols=len(header))
        table.style = 'Table Grid'
        insert_header(table, header)

        row_idx = 0

        for word, word_data in data.items():
            row_idx += 1
            row = table.rows[row_idx]
            table_builder = (
                # FIXME: fix idx here! calculate it because if 'syn' does not exist
                # FIXME: it throws tuple IndexError
                ("syn", "synonyms", settings.user["output"]["synonyms_count"]),
                ("ant", "antonyms", settings.user["output"]["antonyms_count"]),
                ("def", "definitions", None),
            )

            row.cells[0].text = str(row_idx)
            row.cells[1].text = word

            for lexical_category in word_data["translation"]["lexicalCategories"]:
                insert_lexical_category(row.cells[2], lexical_category, "translations")

            for collocation in word_data["translation"]["collocations"]:
                if row.cells[2].paragraphs[0].text != "":
                    row.cells[2].add_paragraph()

                row.cells[2].paragraphs[len(row.cells[2].paragraphs) - 1].text = "- " + collocation

            for lexical_category in word_data["api"].get("lexicalEntries", []):
                col_idx = 3

                for settings_key, cat_key, settings_constraint in table_builder:
                    if self.columns_settings.get(settings_key, False):
                        insert_lexical_category(row.cells[col_idx], lexical_category, cat_key, settings_constraint)
                        col_idx += 1

        document.save(self.filename)

        return True
